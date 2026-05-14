from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Estilos base ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x52, 0xb7, 0x88)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def nota(text):
    p = doc.add_paragraph()
    run = p.add_run(f'[NOTA: {text}]')
    run.italic = True
    run.font.color.rgb = RGBColor(0xaa, 0x66, 0x00)

def separador():
    doc.add_paragraph('─' * 60)

# ══════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════
doc.add_picture  # no imagen, solo texto
t = doc.add_heading('III CONGRESO PERUANO DE PRIMATOLOGÍA 2026', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Integrando Ciencias, Saberes y Sociedad para la Conservación en el Perú')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

doc.add_paragraph('30 de julio al 02 de agosto de 2026').alignment == WD_ALIGN_PARAGRAPH.CENTER
body('Sede: Facultad de Medicina Veterinaria de la UNMSM (Av. Circunvalación 28, San Borja, Lima)')
body('Correo: primatesperuapp@gmail.com')
body('Organizador: Asociación Peruana de Primatología (APP)')
nota('Documento generado para revisión interna. Por favor edita directamente en este archivo.')
doc.add_page_break()

# ══════════════════════════════════════════════
# 1. SOBRE EL CONGRESO
# ══════════════════════════════════════════════
h1('1. Sobre el Congreso')
h2('Presentación')
body('El III Congreso Peruano de Primatología 2026 es el evento nacional dedicado al estudio, conservación y conocimiento de los primates del Perú. Nace de la necesidad urgente de articular a la comunidad científica, instituciones gubernamentales y comunidades locales en torno a los primates neotropicales.')
body('El territorio peruano alberga una de las diversidades de primates más extraordinarias del planeta. Con aproximadamente 47 especies reconocidas (55 taxones), el Perú se posiciona en la vanguardia de la biodiversidad. Resulta fundamental destacar que 11 de estos taxones son estrictamente endémicos, lo que confiere al Estado y la sociedad peruana una responsabilidad ineludible en la custodia global de este patrimonio genético.')
body('Organizado en alianza con universidades nacionales, organismos de gobierno y socios internacionales, el congreso es un espacio de debate científico riguroso y generación de compromisos concretos para la conservación participativa.')

h2('El Perú en números')
tabla = doc.add_table(rows=1, cols=2)
tabla.style = 'Table Grid'
tabla.rows[0].cells[0].text = 'Dato'
tabla.rows[0].cells[1].text = 'Descripción'
datos = [
    ('47 especies', 'de primates registradas'),
    ('11 taxones', 'estrictamente endémicos'),
    ('15 géneros', 'representados en 5 familias'),
    ('4° lugar', 'mundial en diversidad de primates'),
]
for num, label in datos:
    row = tabla.add_row()
    row.cells[0].text = num
    row.cells[1].text = label

h2('Contexto de Conservación')
body('Más allá de su valor intrínseco y evolutivo, los primates desempeñan un rol ecológico, cultural y social insustituible. Como dispersores de semillas y polinizadores, son críticos para la regeneración natural de los bosques y su capacidad para actuar como sumideros de carbono frente al cambio climático. Culturalmente, están profundamente entrelazados con las cosmovisiones y medios de vida de los pueblos indígenas amazónicos y andinos.')
body('A pesar de su importancia, presiones antropogénicas como la expansión agrícola, la minería ilegal y el tráfico han precipitado una crisis. Como respuesta, el Estado peruano, a través de SERFOR, promulgó el Plan Nacional de Conservación de los Primates Amenazados (2019–2029), un documento rector para revertir este declive.')

h3('Alineación Estratégica con SERFOR')
items = [
    'Recuperación de Hábitats: Promoción de corredores ecológicos, modelamiento espacial y restauración de áreas degradadas (Línea de Acción 1).',
    'Reducción de Tráfico y Caza: Análisis de rutas de tráfico ilegal, rescate, protocolos veterinarios y acuerdos comunitarios de subsistencia (Líneas 2 y 4).',
    'Información Consistente: Levantamiento masivo de datos poblacionales, genómicos y taxonómicos para alimentar el SNIFFS estatal (Línea 5).',
    'Gobernanza Participativa: Fortalecimiento de capacidades técnicas con Gobiernos Regionales (GORE), ARFFS y organizaciones indígenas (Líneas 6, 7 y 8).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

h2('Comité Organizador')
nota('Los nombres del comité serán anunciados próximamente. Editar esta sección con los nombres reales.')
comite = [
    ('Presidenta', 'Por confirmar', 'Asociación Peruana de Primatología (APP)'),
    ('Secretaría Científica', 'Por confirmar', 'UNMSM / Universidad Peruana'),
    ('Comité de Programa', 'Por confirmar', 'Instituciones aliadas'),
    ('Comité de Comunicaciones', 'Por confirmar', 'Equipo CPP 2026'),
]
tc = doc.add_table(rows=1, cols=3)
tc.style = 'Table Grid'
tc.rows[0].cells[0].text = 'Rol'
tc.rows[0].cells[1].text = 'Nombre'
tc.rows[0].cells[2].text = 'Institución'
for rol, nombre, inst in comite:
    row = tc.add_row()
    row.cells[0].text = rol
    row.cells[1].text = nombre
    row.cells[2].text = inst

doc.add_page_break()

# ══════════════════════════════════════════════
# 2. PONENTES
# ══════════════════════════════════════════════
h1('2. Ponentes')
h2('Conferenciantes Magistrales')
magistrales = [
    ('Rosamira Guillén Monroy', 'Directora Ejecutiva', 'Fundación Proyecto Tití', 'Colombia', 'Conservación de primates', 'Arquitecta paisajista y conservacionista con más de 25 años de experiencia en la protección del tití cabeciblanco. Lidera proyectos de restauración de ecosistemas y desarrollo comunitario en el Caribe colombiano. Ganadora del Whitley Award y el National Geographic Society/Buffett Award.'),
    ('Dr. Filippo Aureli', 'Investigador Titular C', 'Universidad Veracruzana / Instituto de Neuroetología', 'México', 'Etología y socioecología de primates', 'Etólogo con más de 40 años de experiencia estudiando el comportamiento social, manejo de conflictos y dinámicas de fisión-fusión en primates. Especialista en la conservación del mono araña en México y Centroamérica. Explorador de National Geographic y autor de más de 190 artículos científicos.'),
]
for nombre, titulo, afil, pais, area, bio in magistrales:
    h3(nombre)
    body(f'Cargo: {titulo} | {afil} ({pais})')
    body(f'Área: {area}')
    body(f'Biografía: {bio}')
    doc.add_paragraph()

h2('Ponentes Invitados')
invitados = [
    ('Dr. Alejandro García-Naranjo', 'Director de Conservación', 'Wildlife Conservation Society Perú', 'Colombia', 'Conservación in situ', 'Lidera programas de conservación de primates amenazados en la cuenca del Orinoco y la Amazonía occidental. Pionero en el uso de cámaras trampa para el monitoreo de primates nocturnos.'),
    ('Dra. Yamila Roca Pasca', 'Investigadora Asociada', 'UNAM — México', 'México', 'Etnoprimatología', 'Pionera en estudios de coexistencia humano-primate y conocimiento ecológico de comunidades indígenas amazónicas. Su trabajo integra metodologías de ciencias sociales y ecología para entender las relaciones culturales con los primates.'),
    ('Dr. Pablo Stevenson', 'Profesor', 'Universidad de los Andes', 'Colombia', 'Comportamiento y dieta', 'Investigador líder en ecología de comunidades de primates y dispersión de semillas. Coordina el programa de monitoreo de primates a largo plazo en La Macarena, Colombia.'),
    ('Dra. Gisella Orihuela', 'Médico Veterinario Wildlife', 'SERNANP - Perú', 'Perú', 'Medicina de fauna / One Health', 'Especialista en salud de fauna silvestre y zoonosis emergentes. Lidera el programa de vigilancia epidemiológica en primates de áreas naturales protegidas del Perú.'),
]
for nombre, titulo, afil, pais, area, bio in invitados:
    h3(nombre)
    body(f'Cargo: {titulo} | {afil} ({pais})')
    body(f'Área: {area}')
    body(f'Biografía: {bio}')
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════
# 3. INSCRIPCIONES
# ══════════════════════════════════════════════
h1('3. Inscripciones')
nota('Inscripciones Early Bird disponibles hasta el 31 de mayo de 2026.')

h2('Tabla de Precios')
th = doc.add_table(rows=1, cols=7)
th.style = 'Table Grid'
encabezados = ['Categoría', 'Early Bird (APP)', 'Early Bird (General)', 'Regular (APP)', 'Regular (General)', 'En sede (APP)', 'En sede (General)']
for i, enc in enumerate(encabezados):
    th.rows[0].cells[i].text = enc

precios = [
    ('Profesional', '120', '150', '160', '200', '200', '250'),
    ('Estudiante de posgrado', '60', '80', '80', '110', '100', '140'),
    ('Estudiante de pregrado', '30', '40', '40', '55', '50', '70'),
    ('Comunidades / Gestores', '20', '20', '20', '20', '20', '20'),
]
for fila in precios:
    row = th.add_row()
    for i, val in enumerate(fila):
        row.cells[i].text = f'S/ {val}' if i > 0 else val

h2('Datos Bancarios')
bancarios = [
    ('Titular', 'Estación Biológica Kawsay'),
    ('RUC', '20607343609'),
    ('Banco', 'Banco Interbank'),
    ('Tipo de cuenta', 'Cuenta Corriente Soles'),
    ('Cuenta', '335-3003392978'),
    ('CCI', '003-335-003003392978-74'),
    ('Yape', '966 381 468 — Raúl Bello'),
    ('Plin', '966 381 468 — Estación Biológica Kawsay'),
]
tb = doc.add_table(rows=1, cols=2)
tb.style = 'Table Grid'
tb.rows[0].cells[0].text = 'Campo'
tb.rows[0].cells[1].text = 'Dato'
for campo, dato in bancarios:
    row = tb.add_row()
    row.cells[0].text = campo
    row.cells[1].text = dato

h2('Descuento Grupal')
body('20% de descuento adicional por grupos de 3 personas o más. Adjuntar voucher con los nombres de todos los miembros.')
body('Contacto para descuento grupal: primatesperuapp@gmail.com')

doc.add_page_break()

# ══════════════════════════════════════════════
# 4. PROGRAMA PRELIMINAR
# ══════════════════════════════════════════════
h1('4. Programa Preliminar')
nota('Programa sujeto a modificaciones. Las fechas exactas se confirmarán próximamente.')

dias = [
    ('Día 1 — Julio 2026 | Eje: Ciencia y Conocimiento', [
        ('08:00', 'Registro y acreditación de participantes'),
        ('09:00', 'Ceremonia de apertura oficial'),
        ('09:30', 'Charla Magistral 1 — Por confirmar'),
        ('10:30', 'Coffee break'),
        ('11:00', 'Simposio 1 (Sala A) | Simposio 2 (Sala B)'),
        ('12:30', 'Almuerzo'),
        ('14:00', 'Sesión de Ponencias Orales 1'),
        ('15:30', 'Coffee break'),
        ('16:00', 'Sesión de Pósters 1'),
        ('19:00', 'Cena de bienvenida y networking'),
    ]),
    ('Día 2 — Julio 2026 | Eje: Comunidades y Conservación', [
        ('09:00', 'Charla Magistral 2 — Por confirmar'),
        ('10:00', 'Coffee break'),
        ('10:30', 'Simposio 3 (Sala A) | Simposio 4 (Sala B)'),
        ('12:00', 'Almuerzo'),
        ('13:30', 'Mesa Redonda 1'),
        ('15:00', 'Coffee break'),
        ('15:30', 'Sesión de Ponencias Orales 2'),
        ('17:00', 'Actividad cultural'),
    ]),
    ('Día 3 — Julio 2026 | Eje: Acción y Compromiso', [
        ('09:00', 'Talleres Paralelos 1, 2 y 3 (Salas A, B, C)'),
        ('10:00', 'Coffee break'),
        ('10:30', 'Sesión de Ponencias Orales 3'),
        ('12:00', 'Almuerzo'),
        ('13:30', 'Charla Magistral de Cierre — Por confirmar'),
        ('14:30', 'Asamblea de la Asociación Peruana de Primatología (APP)'),
        ('15:30', 'Coffee break'),
        ('16:00', 'Ceremonia de clausura y premiación de mejores trabajos'),
    ]),
]

for titulo_dia, sesiones in dias:
    h2(titulo_dia)
    tp = doc.add_table(rows=1, cols=2)
    tp.style = 'Table Grid'
    tp.rows[0].cells[0].text = 'Hora'
    tp.rows[0].cells[1].text = 'Actividad'
    for hora, actividad in sesiones:
        row = tp.add_row()
        row.cells[0].text = hora
        row.cells[1].text = actividad
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════
# 5. LLAMADO DE RESÚMENES
# ══════════════════════════════════════════════
h1('5. Llamado de Resúmenes')
body('Apertura: 8 de mayo 2026  |  Cierre: 15 de junio 2026')
body('Lugar: Facultad de Medicina Veterinaria UNMSM (Av. Circunvalación 28, San Borja, Lima).')
body('El Comité Organizador del III Congreso Peruano de Primatología se complace en invitar a investigadores, estudiantes, gestores de áreas protegidas, organizaciones indígenas, comunidades locales, centros de rescate y artistas de conservación a presentar sus propuestas.')

h2('Formatos de Presentación')
formatos = [
    ('Ponencias Orales', '15 min', 'Presentación de datos empíricos, resultados de investigaciones biológicas o análisis sociales.'),
    ('Relatos de Experiencia', '15 min', 'Modalidad para gestores y comunidades; narrativa práctica de conservación.'),
    ('Charlas Relámpago', '5 min', 'Formato ágil para investigaciones en curso o innovaciones tecnológicas.'),
    ('Pósteres y Exposiciones', 'Digital / Físico', 'Infografías científicas, mapas participativos o fotografía de conservación.'),
    ('Simposios / Mesas Redondas', '90 min', 'Panel multidisciplinario para discutir temas críticos.'),
    ('Talleres Colaborativos', 'Sesiones', 'Transferencia de habilidades prácticas, técnicas o de políticas.'),
]
tf = doc.add_table(rows=1, cols=3)
tf.style = 'Table Grid'
tf.rows[0].cells[0].text = 'Formato'
tf.rows[0].cells[1].text = 'Duración'
tf.rows[0].cells[2].text = 'Descripción'
for fmt, dur, desc in formatos:
    row = tf.add_row()
    row.cells[0].text = fmt
    row.cells[1].text = dur
    row.cells[2].text = desc

h2('Pautas de Envío')
pautas = [
    'Extensión: máximo 300 palabras (sin título ni autores).',
    'Formato empírico: Contexto/Justificación, Metodología, Resultados, Implicaciones.',
    'Formato de relatos: Situación, Acciones, Desafíos, Impacto.',
    'Lenguaje inclusivo y ética taxonómica requeridos.',
    'Proyectos con pueblos indígenas deben certificar consentimiento informado (CLPI).',
    'Cada autor puede ser primer autor de un máximo de 2 resúmenes.',
    'Se aceptan resúmenes en español e inglés (español es el idioma principal).',
]
for p in pautas:
    doc.add_paragraph(p, style='List Bullet')

h2('Enlace de envío')
body('Formulario: https://forms.gle/Gzcr2ktTmrEnGTTh9')
body('Plantilla Word: /documentos/Modelo_resumen_CPP2026.docx')

doc.add_page_break()

# ══════════════════════════════════════════════
# 6. FECHAS CLAVE
# ══════════════════════════════════════════════
h1('6. Fechas Clave')
fechas = [
    ('08 Mayo 2026', 'Apertura del llamado de resúmenes'),
    ('31 Mayo 2026', 'Fin de inscripciones Early Bird'),
    ('15 Junio 2026', 'Cierre de envío de resúmenes'),
    ('01 Julio 2026', 'Notificación de aceptaciones'),
    ('15 Julio 2026', 'Cierre de inscripciones regulares'),
    ('30 Julio 2026', '🎉 Inicio del III Congreso Peruano de Primatología 2026'),
]
tf2 = doc.add_table(rows=1, cols=2)
tf2.style = 'Table Grid'
tf2.rows[0].cells[0].text = 'Fecha'
tf2.rows[0].cells[1].text = 'Hito'
for fecha, hito in fechas:
    row = tf2.add_row()
    row.cells[0].text = fecha
    row.cells[1].text = hito

doc.add_page_break()

# ══════════════════════════════════════════════
# 7. CONTACTO
# ══════════════════════════════════════════════
h1('7. Contacto')
body('Correo oficial: primatesperuapp@gmail.com')
body('Instagram: https://www.instagram.com/primatesapp')
body('Facebook: https://www.facebook.com/monosperu')
body('Sede: Facultad de Medicina Veterinaria de la UNMSM (Av. Circunvalación 28, San Borja, Lima)')
body('Fecha del evento: 30 de julio al 02 de agosto de 2026')

# Guardar
output = '/home/jota/Documents/APP_WEB_SERVER/CPP2026_Revision_Web.docx'
doc.save(output)
print(f'✅ Documento guardado en: {output}')
